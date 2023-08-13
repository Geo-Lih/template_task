import pandas as pd
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.shared import Pt, RGBColor
from docx.styles.style import _ParagraphStyle
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx import Document
from docx.document import Document
from docx.shared import Cm
from document_setup import document_setup

# YOUR FILE PATH
FILE_PATH = '/home/georgii/Downloads/шаблон заяви.xlsx'


def build_template(grouped_data):
    doc: Document = document_setup()

    for executor_address_value_tuple, value in grouped_data.items():
        executor, address = executor_address_value_tuple
        debtor_count = len(value)

    # ----------------------------------------------------STYLES----------------------------------------------------

    heading_style: _ParagraphStyle = doc.styles['Heading 1']
    heading_style.font.name = 'Times New Roman'
    heading_style.font.size = Pt(12)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    normal_style: _ParagraphStyle = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal_style.paragraph_format.first_line_indent = Cm(1)

    list_number_style = doc.styles['List Number']
    list_number_style.paragraph_format.first_line_indent = Cm(2)

    # ----------------------------------------------------CONTENT----------------------------------------------------

    table = doc.add_table(rows=1, cols=2)
    table.rows[0].height = Cm(6.604)
    table.columns[0].width = Cm(7.3914)
    table.columns[1].width = Cm(9.7536)

    cell: _Cell = table.rows[0].cells[0]
    p: Paragraph = cell.paragraphs[0]
    p.paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run('№ _____________').italic = True
    p.add_run('\nвід «__» ___.2023 р.').italic = True
    cell.add_paragraph('', style='Normal')
    cell.add_paragraph('', style='Normal')
    table_paragraph = cell.add_paragraph('', style='Normal')
    table_paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
    run = table_paragraph.add_run('Стягувач:')
    run.italic = True
    run.underline = True
    cell.add_paragraph('', style='Normal')
    cell.add_paragraph('', style='Normal')
    cell.add_paragraph('', style='Normal')
    cell.add_paragraph('', style='Normal')
    table_paragraph = cell.add_paragraph('', style='Normal')
    table_paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table_paragraph.paragraph_format.first_line_indent = Cm(0)
    run = table_paragraph.add_run('Правонаступник стягувача (заявник):')
    run.italic = True
    run.underline = True

    cell: _Cell = table.rows[0].cells[1]
    p: Paragraph = cell.paragraphs[0]
    p.paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run('Приватний виконавець').bold = True
    p.add_run(f'\n{executor}').bold = True
    p.add_run(f'\n{address}').italic = True
    p.add_run('\n')
    p.add_run('\nТовариство з обмеженою відповідальністю').bold = True
    p.add_run('\n«Вердикт Капітал»').bold = True
    p.add_run('\nКод ЄДРПОУ 36799749')
    p.add_run('\n04053, м. Київ, вул. Кудрявський узвіз, 5-б').italic = True
    p.add_run('\n')
    p.add_run('\nТовариство з обмеженою відповідальністю').bold = True
    p.add_run('\n«Дебт Форс»').bold = True
    p.add_run('\nКод ЄДРПОУ 43577608')
    p.add_run('\n02121, м. Київ, Харківське шосе, б. 201/203 літера ').italic = True
    p.add_run('\n2А, оф. 602').italic = True
    p.add_run('\n')

    doc.add_paragraph('', style='Normal')

    # ----------------------------------------------------STATEMENT----------------------------------------------------

    p = doc.add_paragraph('', style='Normal')
    p.add_run('ЗАЯВА').bold = True
    run = p.add_run('\nпро зупинення вчинення виконавчих дій')
    run.font.bold = True
    run.font.italic = True
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    doc.add_paragraph('', style='Normal')

    p = doc.add_paragraph('', style='Normal')
    p.add_run('У Вас, на примусовому виконанні перебувають виконавчі провадження, вказані '
              'у додатку 1 до цього листа, Стягувачем за якими є ТОВ «Вердикт Капітал».')
    p = doc.add_paragraph('', style='Normal')
    p.add_run('Між ТОВ «ВЕРДИКТ КАПІТАЛ» та ТОВ «КАМПСІС ФІНАНС» було укладено Договори про '
              'відступленні (купівлю-продаж) прав вимоги, відповідно до яких ТОВ «ВЕРДИКТ КАПІТАЛ» '
              'відступило ТОВ «КАМПСІС ФІНАНС» право вимоги до боржників вказаних Додатку 1.')
    p = doc.add_paragraph('', style='Normal')
    p.add_run('Між ТОВ «КАМПСІС ФІНАНС» та ТОВ «ДЕБТ ФОРС» (Заявником) було укладено Договори про '
              'відступлення (купівлю-продаж) прав вимоги, відповідно до якого ТОВ «КАМПСІС ФІНАНС» '
              'відступило ТОВ «ДЕБТ ФОРС», а ТОВ «ДЕБТ ФОРС» набуло право вимоги до боржників вказаних у додатку 1.')
    p = doc.add_paragraph('', style='Normal')
    p.add_run('Станом на теперішній час ТОВ «ДЕБТ ФОРС» звернулося до суду з заявою про заміну сторони '
              'виконавчого провадження. Докази направлення заяв долучаємо в якості додатків до цього листа.')
    p = doc.add_paragraph('', style='Normal')
    p.add_run('Відповідно до ч.5 ст.15 ЗУ «Про виконавче провадження», - У разі вибуття однієї із сторін '
              'виконавець за заявою сторони, а також заінтересована особа мають право звернутися до суду '
              'із заявою про заміну сторони її правонаступником.')
    p = doc.add_paragraph('', style='Normal')
    p.add_run('Відповідно до п.5 ч.1 ст.34 ЗУ «Про виконавче провадження», - Виконавець зупиняє вчинення '
              'виконавчих дій у разі звернення виконавця та/або заінтересованої особи до суду із заявою '
              'про заміну вибулої сторони правонаступником у порядку, встановленому частиною '
              'п’ятою статті 15 цього Закону')
    p = doc.add_paragraph('', style='Normal')
    run = p.add_run(
        'В зв’язку з вищевикладеним та керуючись ст.ст. 15, 34 Закону України «Про виконавче провадження»,-')
    run.font.italic = True

    # ----------------------------------------------------REQUEST----------------------------------------------------

    p = doc.add_paragraph('', style='Normal')
    p.add_run('Прошу:').bold = True
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    doc.add_paragraph('', style='Normal')

    p = doc.add_paragraph('1. ', style='Normal')
    p.add_run('Винести постанову про зупинення вчинення виконавчих дій по '
              'виконавчим провадженням  вказаним у додатку 1.')
    p = doc.add_paragraph('2. ', style='Normal')
    p.add_run('Про результати розгляду даної заяви повідомити заявника за електронною адресою:', )
    doc.add_paragraph('', style='Normal')

    p = doc.add_paragraph(f'Рєєстр боржників та виконавчих проваджень відкритих у ПВ {executor}', style='List Number')
    p.paragraph_format.left_indent = Cm(2)

    p = doc.add_paragraph('Докази відправлення в суд заяв про заміну сторони у виконавчому провадженні.',
                          style='List Number')
    p.paragraph_format.left_indent = Cm(2)

    p = doc.add_paragraph('Витяг з ЄДРПОУ ТОВ «ДЕБТ ФОРС»', style='List Number')
    p.paragraph_format.left_indent = Cm(2)
    doc.add_paragraph('', style='Normal')

    doc = add_director_signature(doc, 'Олександр КУЗЬМЕНКО')
    doc.add_paragraph('', style='Normal')

    p = doc.add_paragraph('', style='Normal')
    p.add_run('Додаток 1 до Заяви про зупинення вчинення виконавчих дій').bold = True
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p = doc.add_paragraph('', style='Normal')
    p.add_run('від «__» ___.2023 р.').bold = True
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    doc.add_paragraph('', style='Normal')

    # ----------------------------------------------------DEBTOR-TABLE--------------------------------------------------

    table = doc.add_table(rows=debtor_count + 1, cols=5)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(0.889)
    table.columns[1].width = Cm(4.902)
    table.columns[2].width = Cm(2.464)
    table.columns[3].width = Cm(6.322)
    table.columns[4].width = Cm(2.870)
    table.rows[0].height = Cm(1.5)
    headers = ['№', 'ПІБ Боржника', '№ АСВП', 'Суд до якого подана заява про заміну сторони у виконавчому провадженні',
               'Дата подачі заяви']

    for col, header in enumerate(headers):
        cell: _Cell = table.rows[0].cells[col]
        p: Paragraph = cell.paragraphs[0]
        p.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.add_run(header).bold = True

    row = 1
    for key, value_list in grouped_data.items():
        for entry in value_list:
            if row >= debtor_count + 1:
                break

            table.rows[row].cells[0].text = str(row)
            table.rows[row].cells[1].text = entry[1]
            table.rows[row].cells[2].text = str(entry[0])
            table.rows[row].cells[3].text = entry[2]
            table.rows[row].cells[4].text = str(entry[3])

            for col in range(5):
                cell = table.rows[row].cells[col]

                p = cell.paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(0)

            row += 1

    doc.add_paragraph('', style='Normal')
    doc.add_paragraph('', style='Normal')
    doc = add_director_signature(doc, 'Олександр КУЗЬМЕНКО')

    return doc


def add_director_signature(doc, name):
    doc.add_paragraph('', style='Normal')
    p = doc.add_paragraph('', style='Normal')
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(12.52))

    director_run = p.add_run('Директор ТОВ «Дебт Форс»')
    director_run.bold = True

    p.add_run('\t')

    name_run = p.add_run(name)
    name_run.bold = True

    return doc


def main():
    with open(FILE_PATH, 'rb') as excel_file:
        df = pd.read_excel(excel_file)
        grouped_data = {}

        for index, row in df.iterrows():
            key = (row['Приватний виконавець'], row['Адреса ПВ'])
            value = (row['№ АСВП'], row['ПІБ Боржника'],
                     row['Суд до якого подана заява про заміну сторони у виконавчому провадженні'],
                     row['Дата подачі заяви'])

            if key not in grouped_data:
                grouped_data[key] = []

            grouped_data[key].append(value)

    if len(grouped_data) == 0:
        raise Exception('Файл порожній. Немає даних для обробки.')

    for executor_key, grouped_data_for_executor in grouped_data.items():
        doc = build_template({executor_key: grouped_data_for_executor})
        executor_name = executor_key[0]
        # YOUR FILE PATH
        output_file_path = f'/home/georgii/Downloads/{executor_name}.docx'
        doc.save(output_file_path)
        print(f"Документ для '{executor_name}' збережено у файл '{output_file_path}'")


if __name__ == '__main__':
    main()
